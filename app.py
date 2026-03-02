import streamlit as st
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
import io
from docx import Document

# ============================================================
# ESTRUTURAS DE DADOS (SCHEMA PARA A IA) — OPÇÃO 1 (SEM dict)
# ============================================================

class RotinaItem(BaseModel):
    inicio: str = Field(default="00:00")
    fim: str = Field(default="00:00")
    titulo: str = Field(default="Atividade")
    descricao: str = Field(default="Sem descrição")

class AtividadeItem(BaseModel):
    momento: str = Field(default="Momento")
    titulo: str = Field(default="Título")
    descricao: str = Field(default="Descrição")
    mediacao: str = Field(default="Mediação")
    observacao: str = Field(default="Observação")

class CabecalhoPlano(BaseModel):
    tipo_plano: str = Field(default="Diário")
    turma: str = Field(default="Pré I")
    data: str = Field(default="00/00/0000")
    faixa_etaria: str = Field(default="4 anos")
    tema: str = Field(default="(sem tema)")

class PlanoGerado(BaseModel):
    cabecalho: CabecalhoPlano
    campos_experiencia: list[str] = Field(default_factory=list)
    objetivos: list[str] = Field(default_factory=list)
    rotina: list[RotinaItem] = Field(default_factory=list)
    materiais: list[str] = Field(default_factory=list)
    atividades: list[AtividadeItem] = Field(default_factory=list)
    avaliacao: list[str] = Field(default_factory=list)
    adaptacoes: list[str] = Field(default_factory=list)
    observacoes: str = Field(default="(sem observações)")

# ============================================================
# FUNÇÕES AUXILIARES
# ============================================================

def plan_to_markdown(p: PlanoGerado) -> str:
    cab = (
        f"**Tipo de plano:** {p.cabecalho.tipo_plano}\n"
        f"**Turma/Grupo:** {p.cabecalho.turma}\n"
        f"**Data:** {p.cabecalho.data}\n"
        f"**Faixa etária:** {p.cabecalho.faixa_etaria}\n"
        f"**Tema:** {p.cabecalho.tema}"
    )

    rotina_md = "\n".join(
        [f"- **{r.inicio}–{r.fim} | {r.titulo}:** {r.descricao}" for r in p.rotina]
    ) if p.rotina else "- (sem rotina)"

    objetivos_md = "\n".join([f"- {x}" for x in p.objetivos]) if p.objetivos else "- (sem objetivos)"
    campos_md = "\n".join([f"- {x}" for x in p.campos_experiencia]) if p.campos_experiencia else "- (sem campos)"
    materiais_md = "\n".join([f"- {x}" for x in p.materiais]) if p.materiais else "- (sem materiais)"
    avaliacao_md = "\n".join([f"- {x}" for x in p.avaliacao]) if p.avaliacao else "- (sem avaliação)"
    adapt_md = "\n".join([f"- {x}" for x in p.adaptacoes]) if p.adaptacoes else "- (sem adaptações)"

    atv_md_parts = []
    for a in p.atividades:
        atv_md_parts.append(
            f"### {a.momento}\n"
            f"**Atividade:** {a.titulo}\n\n"
            f"**Como fazer:** {a.descricao}\n\n"
            f"**Mediação do professor:** {a.mediacao}\n\n"
            f"**O que observar:** {a.observacao}\n"
        )
    atv_md = "\n".join(atv_md_parts) if atv_md_parts else "### (sem atividades descritas)"

    return f"""# Plano de Aula – Educação Infantil

{cab}

## Campos de Experiência (BNCC)
{campos_md}

## Objetivos
{objetivos_md}

## Rotina
{rotina_md}

## Materiais
{materiais_md}

## Sequência de atividades (descritas)
{atv_md}

## Avaliação (observação e registros)
{avaliacao_md}

## Adaptações / Inclusão
{adapt_md}

## Observações
{p.observacoes}
"""

def create_docx(md_text: str) -> bytes:
    doc = Document()
    doc.add_heading("Plano de Aula – Educação Infantil", 0)

    for line in md_text.split("\n"):
        if line.startswith("# "):
            continue
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line.replace("- ", ""), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ============================================================
# INTERFACE STREAMLIT
# ============================================================

st.set_page_config(page_title="Gerador de Planos de Aula", page_icon="✨", layout="wide")

st.title("✨ Gerador de Planos de Aula")
st.subheader("Educação Infantil com IA")

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.markdown("### 1. Informações Básicas")
    c1, c2 = st.columns(2)
    with c1:
        tipo = st.selectbox("Tipo de plano", ["Diário", "Semanal"])
        turma = st.text_input("Turma/Grupo", value="Pré I")
    with c2:
        data_plano = st.date_input("Data")
        faixa = st.selectbox("Faixa etária", ["3 anos", "4 anos", "5 anos"])

    st.markdown("### 2. Diretrizes para a IA")
    prompt = st.text_area(
        "O que você quer trabalhar? (Tema, foco, ideias)",
        height=150,
        placeholder="Ex: Quero trabalhar a coordenação motora fina usando elementos da natureza..."
    )

    uploaded_files = st.file_uploader(
        "Documentos de Referência (Opcional)",
        accept_multiple_files=True,
        type=["pdf", "txt", "png", "jpg", "jpeg"]
    )

with col2:
    st.markdown("### Prévia e exportação")

    if st.button("✨ Gerar plano com IA", use_container_width=True, type="primary"):
        if "GEMINI_API_KEY" not in st.secrets:
            st.error("⚠️ Chave da API do Gemini não encontrada. Configure os Secrets no Streamlit.")
            st.stop()

        client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])

        with st.spinner("Gerando com IA..."):
            try:
                data_str = data_plano.strftime("%d/%m/%Y")

                prompt_text = f"""Você é um especialista em Educação Infantil e na BNCC.
Crie um plano de aula detalhado para Educação Infantil.

IMPORTANTE:
- Retorne APENAS um JSON válido.
- O JSON deve seguir ESTRITAMENTE o schema.
- No campo "cabecalho", preencha: tipo_plano, turma, data, faixa_etaria, tema.

Parâmetros:
- Faixa etária: {faixa}
- Turma: {turma}
- Data: {data_str}
- Tipo de plano: {tipo}

Diretrizes do usuário:
{prompt if prompt else "Sugira um tema e foco adequados para a faixa etária."}

Regras:
- Preencha todos os campos com riqueza de detalhes.
- Capriche na mediação do professor e no que observar (avaliação).
"""

                contents: list = [prompt_text]

                # Adiciona os arquivos anexados (se existirem)
                if uploaded_files:
                    for f in uploaded_files:
                        contents.append(
                            types.Part.from_bytes(data=f.getvalue(), mime_type=f.type)
                        )

                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=contents,
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        response_schema=PlanoGerado,
                        temperature=0.7,
                    ),
                )

                plano: PlanoGerado = response.parsed
                md_text = plan_to_markdown(plano)

                st.success("Plano gerado com sucesso!")

                with st.container(height=400):
                    st.markdown(md_text)

                d1, d2 = st.columns(2)
                with d1:
                    st.download_button(
                        label="📄 Baixar .txt",
                        data=md_text,
                        file_name=f"plano_{data_str.replace('/','')}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                with d2:
                    docx_data = create_docx(md_text)
                    st.download_button(
                        label="📝 Baixar .docx",
                        data=docx_data,
                        file_name=f"plano_{data_str.replace('/','')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

            except Exception as e:
                st.error(f"Ocorreu um erro ao gerar o plano: {str(e)}")
